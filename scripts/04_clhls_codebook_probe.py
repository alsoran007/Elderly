"""Read-only CLHLS codebook and longitudinal-file probe."""

from pathlib import Path
import csv
import re
import sys


PROJECT_ROOT = Path.cwd()
RAW_ROOT = (PROJECT_ROOT / ".." / "sql").resolve()
CODEBOOK_ROOT = RAW_ROOT / "CLHLS" / "CLHLS_codebook 1998-2018"
DATA_ROOT = RAW_ROOT / "CLHLS"
OUT_ROOT = PROJECT_ROOT / "results" / "data_audit"
OUT_ROOT.mkdir(parents=True, exist_ok=True)

try:
    from docx import Document
except ImportError as exc:
    raise SystemExit("python-docx is required for this probe") from exc


PATTERN = re.compile(
    r"death|died|deceas|mort|surviv|life status|date|year|month|day|"
    r"identifier|identification|\bpid\b|\bid\b|\bhhid\b|age|birth",
    re.IGNORECASE,
)


def iter_docx_text(path: Path):
    doc = Document(path)
    for idx, paragraph in enumerate(doc.paragraphs, start=1):
        text = " ".join(paragraph.text.split())
        if text:
            yield "paragraph", idx, text
    for table_idx, table in enumerate(doc.tables, start=1):
        for row_idx, row in enumerate(table.rows, start=1):
            cells = [" ".join(cell.text.split()) for cell in row.cells]
            text = " | ".join(cell for cell in cells if cell)
            if text:
                yield f"table_{table_idx}_row", row_idx, text


def write_csv(path: Path, rows, fieldnames):
    with path.open("w", newline="", encoding="utf-8-sig") as handle:
        writer = csv.DictWriter(handle, fieldnames=fieldnames)
        writer.writeheader()
        writer.writerows(rows)


def main():
    candidate_rows = []
    document_summary = []
    for docx_path in sorted(CODEBOOK_ROOT.glob("*.docx")):
        all_count = 0
        hit_count = 0
        seen = set()
        for element, index, text in iter_docx_text(docx_path):
            all_count += 1
            if not PATTERN.search(text):
                continue
            key = (element, text)
            if key in seen:
                continue
            seen.add(key)
            hit_count += 1
            candidate_rows.append(
                {
                    "source_doc": str(docx_path.relative_to(RAW_ROOT)).replace("\\", "/"),
                    "element": element,
                    "index": index,
                    "text": text,
                }
            )
        document_summary.append(
            {
                "source_doc": str(docx_path.relative_to(RAW_ROOT)).replace("\\", "/"),
                "text_elements": all_count,
                "candidate_elements": hit_count,
            }
        )

    write_csv(
        OUT_ROOT / "clhls_codebook_candidates.csv",
        candidate_rows,
        ["source_doc", "element", "index", "text"],
    )
    write_csv(
        OUT_ROOT / "clhls_codebook_document_summary.csv",
        document_summary,
        ["source_doc", "text_elements", "candidate_elements"],
    )

    # Import only metadata from each available longitudinal SPSS file when
    # pyreadstat is available. This does not materialize the data in memory.
    metadata_rows = []
    try:
        import pyreadstat
    except ImportError:
        pyreadstat = None

    sav_paths = sorted(DATA_ROOT.rglob("*.sav"))
    for sav_path in sav_paths:
        row = {
            "source_file": str(sav_path.relative_to(RAW_ROOT)).replace("\\", "/"),
            "exists": True,
            "read_status": "not_attempted",
            "rows": "",
            "columns": "",
            "candidate_variables": "",
            "read_error": "",
        }
        if pyreadstat is None:
            row["read_status"] = "pyreadstat_unavailable"
            metadata_rows.append(row)
            continue
        try:
            _, meta = pyreadstat.read_sav(str(sav_path), metadataonly=True)
            names = list(meta.column_names)
            labels = list(meta.column_labels or [""] * len(names))
            hits = [
                name
                for name, label in zip(names, labels)
                if PATTERN.search(f"{name} {label}")
            ]
            row.update(
                {
                    "read_status": "metadata_ok",
                    "rows": meta.number_rows,
                    "columns": len(names),
                    "candidate_variables": ";".join(hits),
                }
            )
        except Exception as exc:  # report file-specific metadata failures
            row["read_status"] = "error"
            row["read_error"] = str(exc)
        metadata_rows.append(row)

    write_csv(
        OUT_ROOT / "clhls_longitudinal_metadata_candidates.csv",
        metadata_rows,
        [
            "source_file",
            "exists",
            "read_status",
            "rows",
            "columns",
            "candidate_variables",
            "read_error",
        ],
    )

    lines = [
        "# CLHLS Codebook Probe",
        "",
        "This probe reads CLHLS codebook DOCX files and SPSS metadata only. It does not modify raw data, recode variables, or construct the final outcome.",
        "",
        f"- Codebook documents found: {len(document_summary)}",
        f"- Candidate codebook text elements: {len(candidate_rows)}",
        f"- Longitudinal SPSS files found: {len(sav_paths)}",
        f"- Metadata reads successful: {sum(row['read_status'] == 'metadata_ok' for row in metadata_rows)}",
        f"- Metadata reads with errors: {sum(row['read_status'] == 'error' for row in metadata_rows)}",
        "",
        "Candidate text still requires manual codebook confirmation. In particular, distinguish death date, death status, interview date, and birth date fields before outcome construction.",
        "",
        "Outputs:",
        "- clhls_codebook_candidates.csv",
        "- clhls_codebook_document_summary.csv",
        "- clhls_longitudinal_metadata_candidates.csv",
    ]
    (OUT_ROOT / "clhls_codebook_probe_summary.md").write_text(
        "\n".join(lines) + "\n", encoding="utf-8"
    )
    print(f"CLHLS codebook probe completed. Outputs: {OUT_ROOT}")


if __name__ == "__main__":
    main()
