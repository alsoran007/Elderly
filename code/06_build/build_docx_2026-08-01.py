#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Build a submission-format DOCX from the Markdown manuscript, placing each figure
immediately after the paragraph that first cites it.

No pandoc dependency: parses the subset of Markdown actually used in these
manuscripts (headings, paragraphs, pipe tables, bold/italic/code inline spans)
and writes DOCX directly via python-docx.

Usage:
    python code/06_build/build_docx_2026-08-01.py EN
    python code/06_build/build_docx_2026-08-01.py CN
"""
import sys, os, re
from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

ROOT = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
FIGDIR = os.path.join(ROOT, "build", "figures")

LANG = (sys.argv[1] if len(sys.argv) > 1 else "EN").upper()
SRC = {
    "EN": "docs/manuscript_full_EN_2026-07-31.md",
    "CN": "docs/manuscript_full_CN_2026-07-31.md",
}[LANG]
OUT = os.path.join(ROOT, "build", f"manuscript_{LANG}_2026-08-01.docx")

# Figure key -> (file, caption). Captions kept short; full legends stay in the
# figure list at the end of the manuscript.
FIGSPEC = [
    ("Figure 1",  "Figure1.png"),
    ("Figure 2",  "Figure2.png"),
    ("Figure 3",  "Figure3.png"),
    ("Figure 4",  "Figure4.png"),
    ("Figure 5",  "Figure5.png"),
    ("Figure S1", "FigureS1.png"),
    ("Figure S2", "FigureS2.png"),
    ("Figure S3", "FigureS3.png"),
    ("Figure S4", "FigureS4.png"),
]
# CN manuscripts cite figures as 图 1 / 图 S1
if LANG == "CN":
    FIGSPEC = [(k.replace("Figure S", "图 S").replace("Figure ", "图 "), v)
               for k, v in FIGSPEC]

BODY_FONT = "Times New Roman" if LANG == "EN" else "SimSun"
MONO_FONT = "Consolas"


def setup(doc):
    s = doc.sections[0]
    s.page_width, s.page_height = Mm(210), Mm(297)
    for m in ("top", "bottom", "left", "right"):
        setattr(s, f"{m}_margin", Mm(25))
    st = doc.styles["Normal"]
    st.font.name = BODY_FONT
    st.font.size = Pt(11 if LANG == "EN" else 10.5)
    st.paragraph_format.line_spacing = 2.0          # double-spaced for review
    st.paragraph_format.space_after = Pt(0)
    # East Asian font binding so CJK renders with the intended face
    st.element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")


INLINE = re.compile(r"(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+?`)")


def add_runs(p, text):
    """Render bold / italic / code inline spans."""
    for tok in INLINE.split(text):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            r = p.add_run(tok[2:-2]); r.bold = True
        elif tok.startswith("*") and tok.endswith("*"):
            r = p.add_run(tok[1:-1]); r.italic = True
        elif tok.startswith("`") and tok.endswith("`"):
            r = p.add_run(tok[1:-1]); r.font.name = MONO_FONT; r.font.size = Pt(9.5)
        else:
            p.add_run(tok)


def add_table(doc, rows):
    """rows: list of cell-lists; first row is the header."""
    ncol = max(len(r) for r in rows)
    t = doc.add_table(rows=0, cols=ncol)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(rows):
        cells = t.add_row().cells
        for j in range(ncol):
            txt = row[j] if j < len(row) else ""
            cell = cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_after = Pt(0)
            add_runs(p, txt)
            for r in p.runs:
                r.font.size = Pt(8.5)
                if i == 0:
                    r.bold = True
    return t


def insert_figure(doc, key, fname):
    path = os.path.join(FIGDIR, fname)
    if not os.path.exists(path):
        print(f"  !! missing {fname}")
        return False
    doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.0
    # 165 mm keeps a safety margin inside the 160 mm text column at A4/25 mm
    p.add_run().add_picture(path, width=Mm(160))
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.line_spacing = 1.0
    cap.paragraph_format.space_after = Pt(6)
    r = cap.add_run(f"[{key}]")
    r.bold = True; r.font.size = Pt(9)
    doc.add_paragraph()
    return True


def main():
    md = open(os.path.join(ROOT, SRC), encoding="utf-8").read().splitlines()
    doc = Document()
    setup(doc)

    pending = list(FIGSPEC)      # figures not yet placed
    placed = []
    tbuf = []                    # pipe-table buffer

    def flush_table():
        if not tbuf:
            return
        rows = []
        for ln in tbuf:
            if re.match(r"^\|[\s:\-|]+\|$", ln):     # separator row
                continue
            rows.append([c.strip() for c in ln.strip().strip("|").split("|")])
        if rows:
            add_table(doc, rows)
            doc.add_paragraph().paragraph_format.space_after = Pt(4)
        tbuf.clear()

    for raw in md:
        ln = raw.rstrip()

        if ln.startswith("|"):
            tbuf.append(ln); continue
        flush_table()

        if not ln.strip():
            continue

        if ln.startswith("# "):
            h = doc.add_heading(ln[2:].strip(), level=0)
            for r in h.runs:
                r.font.name = BODY_FONT; r.font.color.rgb = RGBColor(0, 0, 0)
            continue
        if ln.startswith("### "):
            h = doc.add_heading(ln[4:].strip(), level=3)
            for r in h.runs:
                r.font.name = BODY_FONT; r.font.color.rgb = RGBColor(0, 0, 0)
            continue
        if ln.startswith("## "):
            title = ln[3:].strip()
            h = doc.add_heading(title, level=2)
            for r in h.runs:
                r.font.name = BODY_FONT; r.font.color.rgb = RGBColor(0, 0, 0)
            continue
        if ln.startswith("---"):
            continue
        if ln.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, ln[2:].strip())
            continue

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        add_runs(p, ln)

        # After writing this paragraph, drop in any figure it is the first to cite
        for key, fname in list(pending):
            if re.search(re.escape(key) + r"(?![0-9])", ln):
                if insert_figure(doc, key, fname):
                    placed.append(key)
                pending.remove((key, fname))

    flush_table()

    # Anything never cited in body text goes at the end so nothing is dropped
    if pending:
        doc.add_page_break()
        h = doc.add_heading("Figures not cited in body text" if LANG == "EN"
                            else "正文未引用的图", level=2)
        for r in h.runs:
            r.font.name = BODY_FONT; r.font.color.rgb = RGBColor(0, 0, 0)
        for key, fname in pending:
            insert_figure(doc, key, fname)

    os.makedirs(os.path.dirname(OUT), exist_ok=True)
    doc.save(OUT)
    print(f"saved {OUT}")
    print(f"  placed inline ({len(placed)}): {', '.join(placed)}")
    print(f"  appended at end ({len(pending)}): "
          f"{', '.join(k for k, _ in pending) if pending else 'none'}")


if __name__ == "__main__":
    main()
